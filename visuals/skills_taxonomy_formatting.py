import pandas as pd
import json
from anytree import Node, RenderTree, AsciiStyle
from docx import Document
from docx.shared import Pt
import re


def get_taxonomy_path_description(taxonomy_id, taxonomy_dict):
    """
    Given a taxonomy ID (e.g., "6.4.2.2"), traverse the JSON tree to extract the full path with descriptions.
    """
    parts = taxonomy_id.split('.')
    current_node = taxonomy_dict
    descriptions = []
    
    for part in parts:
        if part in current_node:
            description = current_node[part].get('_description', f"Missing description for {part}")
            descriptions.append(f"{part} {description}")
            current_node = current_node[part]  # Move deeper into the tree
        else:
            return f"Invalid Taxonomy ID: {taxonomy_id}"

    return descriptions

def build_taxonomy_tree_from_path(path_descriptions):
    """
    Given a list of descriptions (e.g., ["6 High-Pressure Hydrogen Gas Systems", ...]),
    build a tree structure using anytree.
    """
    root = None
    current_node = None

    for level, desc in enumerate(path_descriptions):
        if level == 0:
            root = Node(desc)
            current_node = root
        else:
            current_node = Node(desc, parent=current_node)

    return root

def map_skills_to_taxonomy_location(skill_mapping_df, taxonomy_dict):
    """
    Map each skill to its taxonomy path as formatted text.
    """
    def get_formatted_taxonomy_path(taxonomy_id):
        # Get the path descriptions from the taxonomy dictionary
        path_descriptions = get_taxonomy_path_description(taxonomy_id, taxonomy_dict)
        if isinstance(path_descriptions, str):  # Handle errors
            return path_descriptions
        
        # Build the minimal tree for this path
        root = build_taxonomy_tree_from_path(path_descriptions)
        
        # Return the formatted taxonomy tree path
        return "\n".join([f"{pre}{node.name}" for pre, _, node in RenderTree(root, style=AsciiStyle())])

    # Add the taxonomy path to the DataFrame
    skill_mapping_df['taxonomy_location'] = skill_mapping_df['Taxonomy ID'].apply(get_formatted_taxonomy_path)
    return skill_mapping_df

def generate_chapter_format(skill_mapping_df):

    # Create a five-digit sort key where each level has two digits, ensuring consistent sorting
    skill_mapping_df['sort_key'] = skill_mapping_df['Taxonomy ID'].apply(
        lambda x: ''.join([f'{int(part):02d}' for part in x.split('.')])
    )

    #skill_mapping_df['']

    # Debugging step to verify sort keys
    #print('\nChecking sort_key:')
    #print(skill_mapping_df.sort_values(by='sort_key')[250:300][['Taxonomy ID', 'sort_key']])

    # Sort by the generated sort key
    skill_mapping_df = skill_mapping_df.sort_values(by='sort_key')

    grouped_skills = skill_mapping_df.groupby('Taxonomy ID', sort=False)
    #specific_group = grouped_skills.get_group('1.1')
    #print(specific_group['taxonomy_location'].unique())

    counter = 0

    #print('\nChecking if groups are sorted:')

    # Iterate through each group in the grouped_skills object
    """for id, group in grouped_skills:
        # Iterate through the rows in the current group
        for _, row in group.iterrows():
            # Print the current element
            print(f'{counter}\t{id}\t{row["sort_key"]}')
            
            # Increment the counter
            counter += 1
            
            # Break if we've printed 300 elements
            if counter == 300:
                break
        
        # Break the outer loop as well if we've printed 300 elements
        if counter == 300:
            break"""

    doc = Document()
    sections = doc.sections
    doc.sections[0].different_first_page_header_footer = True
    footer_txt = ""

    # Define font size for different levels
    font_sizes = {1: 30, 2: 23, 3: 18}

    for index, (id, group) in enumerate(grouped_skills):
        #print(f"\nDebug: Processing group {index + 1} - Taxonomy Location:\n{taxonomy_location}")
        #print(f"Taxonomy IDs in this group: {group['Taxonomy ID'].tolist()}")

        path_levels = 'MAPPING_ERROR'.join(group['taxonomy_location'].unique()).split('\n')
        if (len(group['taxonomy_location'].unique()) > 1):
            print(group['taxonomy_location'].unique())
        path_description = None

        # Iterate over each path level and generate the heading
        for level, path in enumerate(path_levels):
            if level == (len(path_levels) - 1):  # Ensure we're within bounds
                if (level == 0):
                    if index != 0: doc.add_page_break()
                    footer_txt = f'{'MAPPING_ERROR'.join(group['Taxonomy ID'].unique())}: {re.sub(r'^\d+\s+', '', path.replace('+-- ', '').strip())}'
                    #heading = doc.add_heading('New Chapter', 1)
                    #heading.runs[0].font.size = Pt(50)  # Set font size based on depth
                    #doc.add_page_break()

                # Extract the descriptive part of the path, ignoring the number
                taxonomy_id = 'MAPPING_ERROR'.join(group['Taxonomy ID'].unique())
                path_description = re.sub(r'^\d+\s+', '', path.replace('+-- ', '').strip())

                #print(f"Debug: Adding Heading: {taxonomy_id} {path_description}")

                # Add the heading with the appropriate level and font size
                heading = doc.add_heading(f'{' '*(level)}{taxonomy_id}: {path_description}', level=level + 1)
                font_size = font_sizes.get(level + 1, 14)  # Default to 12 if the level exceeds defined sizes
                heading.runs[0].font.size = Pt(font_size)  # Set font size based on depth

        # Add the skills under this taxonomy location
        for _, row in group.iterrows():
            #print(f"Debug: Adding Skill: ID {row['skill_id']} - {row['skill']} (Similarity: {row['Similarity Score']:.4f})")
            paragraph = doc.add_paragraph(f"ID {row['skill_id']}: {row['skill']} (Similarity Score: {row['Similarity Score']:.4f})")
            paragraph.paragraph_format.space_after = Pt(0)

        if not (doc.sections[-1].footer.paragraphs[-1].text == footer_txt):
            doc.sections[-1].footer.paragraphs[-1].clear()
            doc.sections[-1].footer.paragraphs[-1].add_run(footer_txt)
        #footer.add_paragraph(f'{taxonomy_id}: {path_description}')

    doc.save('./mnt/data/skills_in_chapters.docx')
    print("Skills grouped by taxonomy location have been saved as 'skills_in_chapters.docx'.")

def show_skill_taxonomy():
    # Load the taxonomy mapping data
    skill_mapping_df = pd.read_csv('./mnt/data/threshold_skills_insertion.csv')

    # Load the taxonomy structure
    with open('./mnt/data/taxonomy_tree.json', 'r') as f:
        taxonomy_dict = json.load(f)

    # Map skills to the taxonomy tree and add the formatted paths
    updated_skill_mapping_df = map_skills_to_taxonomy_location(skill_mapping_df, taxonomy_dict)

    # Print the updated DataFrame with the taxonomy path
    #print(updated_skill_mapping_df.head())

    updated_skill_mapping_df = updated_skill_mapping_df.drop(columns=['Mapped Node'])

    cols = updated_skill_mapping_df.columns.tolist()

    cols = cols[:3] + [cols[4]] + [cols[3]]

    updated_skill_mapping_df = updated_skill_mapping_df[cols]

    # Optionally save the result to a file
    updated_skill_mapping_df.to_csv('./mnt/data/skill_taxonomy_with_locations.csv', index=False)

    # Generate chapter-like format
    generate_chapter_format(updated_skill_mapping_df)

if __name__ == "__main__":
    show_skill_taxonomy()
